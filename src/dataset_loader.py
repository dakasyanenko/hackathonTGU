"""
Модуль загрузки реального датасета:
- Вакансии из CSV (5_vacancies.csv)
- Резюме из .docx файлов (CV/*.docx)

Датасет: https://github.com/NataliaVanetik/vacancy-resume-matching-dataset

Поддерживает загрузку как с локальной папки, так и автоматическое скачивание
с GitHub (raw content).
"""

from __future__ import annotations

import csv
import re
import urllib.request
from pathlib import Path
from typing import Any

import docx

from src.summarizer import summarize_vacancy


# GitHub raw content URLs
_GITHUB_BASE = "https://raw.githubusercontent.com/NataliaVanetik/vacancy-resume-matching-dataset/refs/heads/main"
_GITHUB_FILES = {
    "csv": f"{_GITHUB_BASE}/5_vacancies.csv",
    "annotations": f"{_GITHUB_BASE}/annotations-for-the-first-30-vacancies.txt",
}

# CV файлы перечислены явно (1.docx .. 65.docx)
_CV_FILES = [f"{i}.docx" for i in range(1, 66)]


def _fetch_text(url: str) -> str:
    """Скачивает текстовый файл с GitHub raw URL."""
    with urllib.request.urlopen(url, timeout=30) as resp:
        return resp.read().decode("utf-8")


def _fetch_bytes(url: str) -> bytes:
    """Скачивает бинарный файл с GitHub raw URL."""
    with urllib.request.urlopen(url, timeout=30) as resp:
        return resp.read()


def _is_github_url(path: str) -> bool:
    """Проверяет, является ли путь URL на GitHub."""
    return path.startswith("http://") or path.startswith("https://")


def load_vacancies_from_csv(csv_source: str) -> list[dict]:
    """
    Загружает вакансии из CSV (файл или URL).

    Args:
        csv_source: Путь к 5_vacancies.csv или GitHub raw URL.

    Returns:
        Список вакансий.
    """
    vacancies = []

    if _is_github_url(csv_source):
        csv_content = _fetch_text(csv_source)
        import io
        reader = csv.DictReader(io.StringIO(csv_content))
        rows = list(reader)
    else:
        with open(csv_source, "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            rows = list(reader)

    for row in rows:
        full_desc = row.get("job_description", "")
        summarized = summarize_vacancy(full_desc)
        skills = _extract_skills_from_job_desc(full_desc)

        vacancies.append({
            "id": row["id"],
            "data": {
                "position": row.get("job_title", ""),
                "description": summarized,
                "industry": "",
                "mandatoryRequirements": _extract_requirements(full_desc),
                "projectTasks": "",
                "experienceLevels": _extract_experience_level(full_desc),
            },
            "skills": skills,
            "dataEng": summarized,
        })

    return vacancies


def _extract_skills_from_job_desc(text: str) -> list[str]:
    """Извлекает технические навыки из описания вакансии."""
    tech_keywords = [
        "Python", "Java", "JavaScript", "TypeScript", "C++", "C#",
        "Ruby", "Go", "Rust", "Kotlin", "Swift", "Scala", "PHP",
        "SQL", "MySQL", "PostgreSQL", "MongoDB", "Redis", "Oracle",
        "React", "Angular", "Vue", "Node.js", "jQuery", "Bootstrap",
        "Docker", "Kubernetes", "AWS", "Azure", "GCP", "Linux", "Windows",
        "Git", "Jenkins", "Jira",
        "REST", "GraphQL", "JSON", "XML", "SOAP", "API",
        "Maven", "Gradle", "JUnit", "PyTest",
        "Spring", "Spring Boot", "Hibernate", "Django", "Flask",
        "LAMP", "Drupal", "ElasticSearch", "Elasticsearch",
        "PHP", "Go", "View.js", "HubSpot",
        "RabbitMQ", "Redis", "EC2", "S3", "Lambda", "CloudFront",
        "MariaDB", "TDD", "Agile", "Scrum",
        "MVC", "OOP", "AOP", "Microservices",
        "ASP.Net", "WCF", "Entity Framework", "Visual Studio", "TFS",
        "MSSQL", ".Net", "jQuery", "Angular",
        "Perl", "Clojure", "Java", "Ruby", "C",
        "tcpdump", "PKI", "SSL", "IPSec", "IPv4", "IPv6",
        "CISSP", "Security+", "DoD 8570",
    ]

    found = []
    text_lower = text.lower()
    for kw in tech_keywords:
        if kw.lower() in text_lower:
            found.append(kw)
    return found


def _extract_requirements(text: str) -> str:
    """Извлекает обязательные требования из текста вакансии."""
    patterns = [
        r"(?:Minimum|Required|Requirements?|Skills)[^:]*:\s*([^.\n]+)",
        r"(?:Must have|Expert in|Proficient with)[^.\n]+",
        r"(?:Bachelor.*?degree|Master.*?degree)[^.\n]*",
        r"(\d+\+?\s+years[^.\n]*)",
    ]
    requirements = []
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        requirements.extend(matches)
    return " ".join(requirements[:5])


def _extract_experience_level(text: str) -> str:
    """Извлекает требуемый уровень опыта."""
    patterns = [
        r"(\d+\+?\s+years[^.\n]*)",
        r"(Senior|Junior|Mid|Mid-level|Lead|Principal|Staff)",
        r"(SME|Journeyman|Entry)",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1)
    return ""


def load_resume_from_docx_source(docx_source: str) -> dict[str, Any]:
    """
    Загружает резюме из .docx (файл или URL).

    Args:
        docx_source: Путь к .docx файлу или GitHub raw URL.

    Returns:
        Словарь резюме.
    """
    if _is_github_url(docx_source):
        data = _fetch_bytes(docx_source)
        import io
        doc = docx.Document(io.BytesIO(data))
    else:
        doc = docx.Document(docx_source)

    # Текст из параграфов
    paragraphs_text = "\n".join(p.text for p in doc.paragraphs)

    # Текст из таблиц
    table_parts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    table_parts.append(cell.text.strip())
    tables_text = " ".join(table_parts)

    full_text = paragraphs_text + "\n" + tables_text

    # Определяем ID
    if _is_github_url(docx_source):
        # Извлекаем имя файла из URL
        resume_id = docx_source.rstrip("/").split("/")[-1].replace(".docx", "")
    else:
        resume_id = Path(docx_source).stem

    return {
        "id": f"cv_{resume_id}",
        "text": full_text,
        "raw_text": full_text,
    }


def load_all_resumes(cv_source: str) -> list[dict[str, Any]]:
    """
    Загружает все резюме (локальная директория или GitHub).

    Args:
        cv_source: Путь к директории CV/ или GitHub raw URL директории.

    Returns:
        Список резюме.
    """
    resumes = []

    if _is_github_url(cv_source):
        # Загружаем с GitHub
        github_cv_base = f"{_GITHUB_BASE}/CV"
        for cv_file in _CV_FILES:
            url = f"{github_cv_base}/{cv_file}"
            try:
                resume = load_resume_from_docx_source(url)
                resumes.append(resume)
            except Exception as e:
                print(f"  ⚠️  Ошибка загрузки {cv_file}: {e}")
    else:
        # Локальная директория
        cv_path = Path(cv_source)
        docx_files = sorted(cv_path.glob("*.docx"))
        for docx_file in docx_files:
            try:
                resume = load_resume_from_docx_source(str(docx_file))
                resumes.append(resume)
            except Exception as e:
                print(f"  ⚠️  Ошибка загрузки {docx_file.name}: {e}")

    print(f"Загружено {len(resumes)} резюме")
    return resumes


# ------------------------------------------------------------------
# Обратная совместимость: старые функции
# ------------------------------------------------------------------

def _to_raw_github_url(url: str) -> str:
    """Преобразует github.com URL в raw.githubusercontent.com URL."""
    url = url.rstrip("/")
    if "raw.githubusercontent.com" in url:
        return url
    # Формат: https://github.com/user/repo или https://github.com/user/repo/tree/branch
    if "/tree/" in url:
        # https://github.com/user/repo/tree/branch -> https://raw.githubusercontent.com/user/repo/branch/
        parts = url.split("/tree/")
        base = parts[0].replace("github.com", "raw.githubusercontent.com")
        return f"{base}/{parts[1]}"
    else:
        # https://github.com/user/repo -> default branch main
        base = url.replace("github.com", "raw.githubusercontent.com")
        return f"{base}/refs/heads/main"


def load_dataset(base_dir: str) -> tuple[list[dict], list[dict]]:
    """
    Загружает весь датасет: вакансии и резюме.
    Поддерживает локальный путь и GitHub URL.

    Args:
        base_dir: Путь к директории датасета ИЛИ GitHub URL репозитория.

    Returns:
        Кортеж (vacancies, resumes)
    """
    if _is_github_url(base_dir):
        raw_base = _to_raw_github_url(base_dir)
        vacancies = load_vacancies_from_csv(f"{raw_base}/5_vacancies.csv")
        resumes = load_all_resumes(f"{raw_base}/CV")
    else:
        base_path = Path(base_dir)

        csv_path = base_path / "5_vacancies.csv"
        if not csv_path.exists():
            raise FileNotFoundError(f"CSV не найден: {csv_path}")
        vacancies = load_vacancies_from_csv(str(csv_path))

        cv_dir = base_path / "CV"
        if not cv_dir.exists():
            raise FileNotFoundError(f"CV директория не найдена: {cv_dir}")
        resumes = load_all_resumes(str(cv_dir))

    print(f"Загружено {len(vacancies)} вакансий, {len(resumes)} резюме")
    return vacancies, resumes
